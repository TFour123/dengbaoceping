from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips
from docx.enum.text import WD_LINE_SPACING
from copy import deepcopy

def set_cell_vertical_merge(cell, merge_type):
    """设置单元格垂直合并"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vMerge = tcPr.find(qn('w:vMerge'))
    if vMerge is None:
        vMerge = OxmlElement('w:vMerge')
        tcPr.append(vMerge)
    vMerge.set(qn('w:val'), merge_type)

def get_cell_style(cell):
    """提取单元格的样式XML元素（边框、背景色等）"""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return {}
    
    style_elements = {}
    for elem_name in ['w:tcBorders', 'w:shd', 'w:tcW', 'w:vAlign']:
        elem = tcPr.find(qn(elem_name))
        if elem is not None:
            style_elements[elem_name] = deepcopy(elem)
    return style_elements

def apply_cell_style(cell, style_elements):
    """将保存的样式应用到单元格"""
    if not style_elements:
        set_default_border(cell)
        return
    
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for elem_name, elem in style_elements.items():
        existing = tcPr.find(qn(elem_name))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(deepcopy(elem))

def set_default_border(cell):
    """设置默认边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ['top', 'left', 'bottom', 'right']:
        tag = 'w:{}'.format(edge)
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:color'), '000000')
        element.set(qn('w:space'), '0')

def set_cell_font_style(cell, font_name='宋体', font_size=Pt(10.5)):
    """
    设置单元格字体样式
    font_name: 字体名称
    font_size: 字号（五号 = 10.5pt）
    行距: 最小值 0磅
    """
    for paragraph in cell.paragraphs:
        # 设置段落行距为最小值 0磅
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        paragraph.paragraph_format.line_spacing = Pt(0)
        # 段前段后间距设为0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            # 设置中文字体（东亚字体）
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def reconstruct_compliance_tables_strict(file_path, output_path):
    doc = Document(file_path)
    print(f"开始处理文件: {file_path}")

    for table in doc.tables:
        if not table.rows:
            continue
            
        header_cells = table.rows[0].cells
        status_col_index = -1
        
        for i, cell in enumerate(header_cells):
            if "符合情况" in cell.text.strip().replace('\n', ''):
                status_col_index = i
                break
        
        if status_col_index == -1:
            continue

        # =======================================================
        # 步骤 1: 提取数据 + 每行的样式（关键改动！）
        # =======================================================
        kept_rows = []  # 存放 (row_data, row_styles) 元组
        current_category = ""
        
        for i in range(1, len(table.rows)):
            row = table.rows[i]
            cells = row.cells
            
            cat_text = cells[0].text.strip()
            if cat_text:
                current_category = cat_text
            
            status_text = cells[status_col_index].text.strip()
            
            # 只保留"不符合"、"部分符合"等
            if status_text not in ["符合", "不适用"]:
                row_data = []
                row_styles = []  # 保存这一行每个单元格的样式
                
                for j, cell in enumerate(cells):
                    # 保存样式
                    row_styles.append(get_cell_style(cell))
                    # 保存文本
                    if j == 0:
                        row_data.append(current_category)
                    else:
                        row_data.append(cell.text.strip())
                
                kept_rows.append((row_data, row_styles))

        # =======================================================
        # 步骤 2: 清空旧行
        # =======================================================
        for i in range(len(table.rows) - 1, 0, -1):
            tbl = table._tbl
            tr = table.rows[i]._element
            tbl.remove(tr)

        # =======================================================
        # 步骤 3: 重写数据，恢复每行原本的样式 + 设置字体
        # =======================================================
        for row_data, row_styles in kept_rows:
            new_row = table.add_row()
            
            for j, text in enumerate(row_data):
                if j < len(new_row.cells):
                    cell = new_row.cells[j]
                    cell.text = text
                    # 应用该行该列原本的样式（边框、背景色）
                    if j < len(row_styles):
                        apply_cell_style(cell, row_styles[j])
                    # 设置字体：宋体、五号、行距最小值0磅
                    set_cell_font_style(cell)

        # =======================================================
        # 步骤 4: 视觉合并第一列
        # =======================================================
        if len(table.rows) > 1:
            current_rows = table.rows
            
            set_cell_vertical_merge(current_rows[1].cells[0], 'restart')
            last_category = current_rows[1].cells[0].text.strip()
            
            for i in range(2, len(current_rows)):
                this_cell = current_rows[i].cells[0]
                this_cat = this_cell.text.strip()
                
                if this_cat == last_category:
                    set_cell_vertical_merge(this_cell, 'continue')
                    this_cell.text = ""
                else:
                    set_cell_vertical_merge(this_cell, 'restart')
                    last_category = this_cat

    doc.save(output_path)
    print(f"处理完成！文件保存为: {output_path}")

# --- 运行配置 ---
input_file = "1.docx"
output_file = "1_cleaned_final.docx"

try:
    reconstruct_compliance_tables_strict(input_file, output_file)
except Exception as e:
    print(f"运行出错: {e}")
